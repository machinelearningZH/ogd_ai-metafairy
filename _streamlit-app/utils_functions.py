import streamlit as st
import re
import pandas as pd
from datetime import datetime
import io
import base64
from docx import Document
from docx.shared import Pt, Inches


def parse_analysis_results(results):
    """Parse the LLM response. Extract the scores and the qualitative analysis.

    Args:
        results (str): The response from the LLM.

    Returns:
        pd.DataFrame: A DataFrame with the scores and the qualitative analysis.

    """
    content = re.findall(r"<dateninhalt>(.*?)</dateninhalt>", results, re.DOTALL)[
        0
    ].strip()
    content_score = re.findall(
        r"<dateninhalt-score>(.*?)</dateninhalt-score>", results, re.DOTALL
    )[0].strip()
    context = re.findall(
        r"<methodik>(.*?)</methodik>",
        results,
        re.DOTALL,
    )[0].strip()
    context_score = re.findall(
        r"<methodik-score>(.*?)</methodik-score>",
        results,
        re.DOTALL,
    )[0].strip()
    quality = re.findall(r"<datenqualität>(.*?)</datenqualität>", results, re.DOTALL)[
        0
    ].strip()
    quality_score = re.findall(
        r"<datenqualität-score>(.*?)</datenqualität-score>",
        results,
        re.DOTALL,
    )[0].strip()
    spacial = re.findall(
        r"<geographie>(.*?)</geographie>",
        results,
        re.DOTALL,
    )[0].strip()
    spacial_score = re.findall(
        r"<geographie-score>(.*?)</geographie-score>",
        results,
        re.DOTALL,
    )[0].strip()

    tmp = pd.DataFrame(
        (
            content,
            content_score,
            context,
            context_score,
            quality,
            quality_score,
            spacial,
            spacial_score,
        )
    ).T
    tmp.columns = [
        "content",
        "content_score",
        "context",
        "context_score",
        "quality",
        "quality_score",
        "spacial",
        "spacial_score",
    ]
    return tmp


def create_download_link(response, time_processed, title=None):
    """Create a downloadable Word document of the result. Create also a download link that can be clicked in the Streamlit app without rerunning the whole app.

    Args:
        response (str): The response from the LLM.
        time_processed (float): The time it took to process the request.
        title (str, optional): The title of the Word document. Defaults to None.
    """
    document = Document()

    h1 = document.add_heading("Datensatzbeschreibung: «" + title + "»")
    p1 = document.add_paragraph("\n" + response)

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der App «Datensätze einfach beschreiben» des Kantons Zürich.\nSprachmodell: OpenAI GPT-4o\nVerarbeitungszeit: {time_processed:.1f} Sekunden"

    # Set font and font size for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Arial"

    # Set font size for all headings.
    for paragraph in [h1]:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    # Set font size for all paragraphs.
    for paragraph in [p1]:
        for run in paragraph.runs:
            run.font.size = Pt(10)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(7)

    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches

    io_stream = io.BytesIO()
    document.save(io_stream)

    # # A download button resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb

    b64 = base64.b64encode(io_stream.getvalue())
    file_name = f"{title}.docx"
    caption = "Datenbeschreibung herunterladen"
    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)
